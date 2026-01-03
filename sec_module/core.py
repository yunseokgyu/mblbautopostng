import requests
import google.generativeai as genai
from bs4 import BeautifulSoup
import re
import time
import os
import json
from dotenv import load_dotenv
import yfinance as yf
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load environment variables
load_dotenv()

# --- Configuration ---
# SEC requires a User-Agent with an email.
SEC_HEADERS = {'User-Agent': "SecAnalyzerBot test@example.com"} 
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    print("Warning: GEMINI_API_KEY not found in environment variables.")

genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-flash-latest')

# --- Data Collection ---

def get_cik_from_ticker(ticker):
    """
    Fetch company_tickers.json to map Ticker -> CIK.
    Returns CIK as a string (padded with zeros if needed).
    """
    try:
        url = "https://www.sec.gov/files/company_tickers.json"
        response = requests.get(url, headers=SEC_HEADERS)
        response.raise_for_status()
        data = response.json()
        
        for key, value in data.items():
            if value['ticker'] == ticker.upper():
                return str(value['cik_str']).zfill(10) # CIK is 10 digits
        print(f"[{ticker}] Ticker not found in SEC database.")
        return None
    except Exception as e:
        print(f"[{ticker}] Error fetching CIK: {e}")
        return None

def get_latest_filing_url(cik, ticker, form_type="10-K"):
    """
    Fetch company submissions to find the latest URL for a specific form type.
    Returns: (url, filing_date)
    """
    try:
        # SEC Submissions API: https://data.sec.gov/submissions/CIK{cik}.json
        url = f"https://data.sec.gov/submissions/CIK{cik}.json"
        response = requests.get(url, headers=SEC_HEADERS)
        response.raise_for_status()
        data = response.json()
        
        filings = data['filings']['recent']
        
        # Iterate to find the latest filing of the requested type
        for i in range(len(filings['accessionNumber'])):
            form = filings['form'][i]
            if form == form_type:
                accession_number = filings['accessionNumber'][i]
                primary_document = filings['primaryDocument'][i]
                filing_date = filings['filingDate'][i]
                
                # Construct URL
                accession_number_no_dashes = accession_number.replace('-', '')
                cik_int = int(cik)
                
                file_url = f"https://www.sec.gov/Archives/edgar/data/{cik_int}/{accession_number_no_dashes}/{primary_document}"
                print(f"[{ticker}] Found {form_type} URL: {file_url} (Date: {filing_date})")
                return file_url, filing_date
                
        print(f"[{ticker}] No {form_type} found in recent submissions.")
        return None, None
    except Exception as e:
        print(f"[{ticker}] Error fetching {form_type} URL: {e}")
        return None, None

def get_sec_data(ticker, form_type="10-K"):
    """
    Orchestrate the download of the Filing HTML.
    Returns: (html_content, filing_date)
    """
    print(f"[{ticker}] Starting data collection for {form_type}...")
    cik = get_cik_from_ticker(ticker)
    if not cik:
        return None, None
    
    url, filing_date = get_latest_filing_url(cik, ticker, form_type)
    if not url:
        return None, None
        
    try:
        response = requests.get(url, headers=SEC_HEADERS)
        response.raise_for_status()
        time.sleep(0.1) 
        return response.text, filing_date
    except Exception as e:
        print(f"[{ticker}] Error downloading HTML: {e}")
        return None, None

def extract_sections(html_content):
    """
    Extract full text from the 10-K HTML, stripping tags.
    """
    print("Preprocessing HTML (Full Text)...")
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Get text and clean up whitespace
    text = soup.get_text(separator='\n')
    text = re.sub(r'\s+', ' ', text).strip() # Collapse multiple spaces/newlines
    
    print(f"Extracted {len(text)} characters.")
    return text

def chunk_text(text, chunk_size=30000):
    """
    Split text into chunks of approximately chunk_size characters.
    """
    return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]

def analyze_with_gemini(text, ticker, filing_date, progress_callback=None, mode="full"):
    """
    Send extracted text to Gemini for translation or summarization.
    mode: 'full' (Detailed Translation) or 'summary' (Executive Summary)
    """
    print(f"Analyzing with Gemini ({mode})...")
    
    chunks = chunk_text(text)
    report_title = "Full Translation" if mode == "full" else "Executive Summary"
    full_report = f"# {ticker} 10-K Report Analysis ({report_title})\n**Filing Date:** {filing_date}\n\n---\n\n"
    
    for i, chunk in enumerate(chunks):
        msg = f"Processing Chunk {i+1}/{len(chunks)} ({len(chunk)} chars)..."
        print(msg)
        if progress_callback:
            progress_callback(i + 1, len(chunks), msg)
        
        if mode == "summary":
            prompt = f"""
            You are a potential power blogger who specializes in US stock analysis.
            Your task is to summarize this part of the SEC 10-K report into an easy-to-read blog post in Korean.

            **Context:** Part {i+1} of {len(chunks)}.
            **Company:** {ticker}
            **Filing Date:** {filing_date}

            **Style Guidelines:**
            - **Tone:** Professional yet accessible (Use polite Korean, "~해요" style).
            - **Formatting:** Clean and professional. **Do NOT use emojis.** Use bolding and bullet points for structure.
            - **Level:** Explain as if talking to a beginner investor. Avoid jargon (or explain it simply).

            **Content Structure:**
            1. **🌟 3-Line Summary:** Start with 3 bullet points summarizing the most important info in this chunk.
            2. **🏢 What does this company do?:** (Only if 'Business' section matches) Briefly explain how they make money.
            3. **📊 Financial Highlights (Must be a Table):**
               - Create a simple Markdown table comparing **Current Q/Y** vs **Previous Q/Y**.
               - Columns: [Metric, Current, Previous, YoY Change].
               - Metrics to include: Revenue, Operating Income, Net Income, EPS.
               - If exact numbers aren't found, use "N/A".
            4. **💰 Shareholder Returns (Dividends & Buybacks):**
               - **Dividends:** Mention current dividend per share and yield if available.
               - **Buybacks (Stock Repurchase):** Detailed amount repurchased and remaining authorization. 
               - **Keyword Check:** Look specifically for "Repurchase Program", "Dividend Declaration".
            5. **🔮 Guidance Check (Important):** If the company provides **Guidance** or **Full Year Outlook**:
               - Create a dedicated section titled "**🔮 가이던스 변경 (Before vs After)**".
               - Explicitly compare the Previous Estimate vs New Estimate.
               - Format: "Old: $X.XX → New: $Y.YY" or similar clearly visible comparison.
               - If no guidance is mentioned, skip this section.
            6. **⚠️ Risk Check:** Highlight any significant risks mentioned.

            **Goal:** Make the user feel smart after reading this. Keep it simple!
            
            **Input Text:**
            """
        else:
            prompt = f"""
            You are a professional translator and investment analyst.
            Your task is to translate the following part of an SEC 10-K report into Korean.
            
            **Context:** Part {i+1} of {len(chunks)}.
            **Company:** {ticker}
            **Filing Date:** {filing_date}
            
            **Instructions:**
            1. **Translate fully and detailedly.** Do not summarize.
            2. **Shareholder Returns (Conditional):** Check this text for **Dividend** or **Buyback** information.
               - **IF FOUND:** Create a dedicated section titled "## Shareholder Returns" and analyze it in detail.
               - **IF NOT FOUND:** Do **NOT** create this section. Skip it to avoid redundancy.
            3. **Maintain original structure.** If the text contains headers (Item 1, etc.), keep them.
            4. **Tone:** Professional, financial.
            5. **Output:** Markdown format.
            
            **Input Text (Part {i+1}):**
            """
        
        retries = 0
        max_retries = 5
        
        while retries < max_retries:
            try:
                response = model.generate_content([prompt, chunk])
                translated_text = response.text
                full_report += translated_text + "\n\n"
                # Rate limit safety (Base sleep)
                time.sleep(5) 
                break
            except Exception as e:
                if "429" in str(e) or "Quota exceeded" in str(e):
                    wait_time = (2 ** retries) * 10 # 10s, 20s, 40s, 80s...
                    print(f"Rate limit hit (Chunk {i+1}). Waiting {wait_time}s...")
                    if progress_callback:
                        progress_callback(i + 1, len(chunks), f"Rate Limit Hit. Waiting {wait_time}s...")
                    time.sleep(wait_time)
                    retries += 1
                else:
                    print(f"Gemini API Error (Chunk {i+1}): {e}")
                    full_report += f"\n\n[Error translating Chunk {i+1}: {e}]\n\n"
                    break
        else:
             full_report += f"\n\n[Failed to translate Chunk {i+1} after retries]\n\n"
            
    return full_report

def get_financials(ticker):
    """
    Fetch financial data using yfinance.
    Returns: (income_stmt, balance_sheet, cash_flow, info)
    """
    print(f"[{ticker}] Fetching financial data from Yahoo Finance...")
    try:
        stock = yf.Ticker(ticker)
        # Get last 3 years
        income = stock.financials.iloc[:, :3]
        balance = stock.balance_sheet.iloc[:, :3]
        cashflow = stock.cashflow.iloc[:, :3]
        info = stock.info
        return income, balance, cashflow, info
    except Exception as e:
        print(f"[{ticker}] Error fetching financials: {e}")
        return None, None, None, None

def set_table_border(table):
    """
    Apply professional borders to the table (Top/Bottom thick).
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # Add borders
    borders = OxmlElement('w:tblBorders')
    
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'thick')
    top.set(qn('w:sz'), '12')
    borders.append(top)
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'thick')
    bottom.set(qn('w:sz'), '12')
    borders.append(bottom)
    
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4')
    borders.append(insideH)
    
    tblPr.append(borders)

def create_key_metrics_table(doc, info):
    """
    Create a summary box for key metrics (Market Cap, PER, PBR, etc.)
    """
    if not info:
        return

    doc.add_heading('Investment Summary', level=2)
    
    # Metrics to display
    metrics = {
        "Current Price": info.get('currentPrice', 'N/A'),
        "Market Cap": info.get('marketCap', 'N/A'),
        "Trailing PER": info.get('trailingPE', 'N/A'),
        "Forward PER": info.get('forwardPE', 'N/A'),
        "PBR": info.get('priceToBook', 'N/A'),
        "ROE": info.get('returnOnEquity', 'N/A'),
        "Dividend Yield": info.get('dividendYield', 'N/A'),
        "Target Price": info.get('targetMeanPrice', 'N/A')
    }
    
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    # Row 1 Headers
    headers = ["Current Price", "Market Cap", "PER (Trailing)", "PBR"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Shading
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'E7E6E6') # Light Gray
        tcPr.append(shd)

    # Row 1 Values
    vals = [metrics["Current Price"], metrics["Market Cap"], metrics["Trailing PER"], metrics["PBR"]]
    for i, v in enumerate(vals):
        cell = table.cell(1, i)
        if isinstance(v, (int, float)):
             # Format large numbers
            if i == 1 and v > 1000000: # Market Cap
                cell.text = f"${v/1000000000:,.2f} B"
            else:
                cell.text = f"{v:,.2f}" if isinstance(v, float) else f"{v:,}"
        else:
            cell.text = str(v)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_table_border(table)
    doc.add_paragraph("\n")

def filter_financial_rows(df, statement_type):
    """
    Select key rows for simplified viewing.
    """
    if df is None or df.empty:
        return df
        
    # Standardize index to string
    df.index = df.index.map(str)
    
    key_rows = []
    if statement_type == "income":
        key_rows = ['Total Revenue', 'Operating Income', 'Net Income', 'Basic EPS', 'Diluted EPS']
    elif statement_type == "balance":
        key_rows = ['Total Assets', 'Total Liabilities Net Minority Interest', 'Stockholders Equity', 'Total Debt', 'Cash And Cash Equivalents']
    elif statement_type == "cashflow":
        key_rows = ['Operating Cash Flow', 'Investing Cash Flow', 'Financing Cash Flow', 'Free Cash Flow']
        
    # Filter rows that exist in the dataframe
    existing_rows = [row for row in key_rows if row in df.index]
    
    if existing_rows:
        return df.loc[existing_rows]
    else:
        # If no key rows found (different naming convention), return top 10 rows
        return df.head(10)

def df_to_word_table(doc, df, title):
    """
    Convert DataFrame to a styled Word table.
    """
    if df is None or df.empty:
        return

    doc.add_heading(title, level=2)
    
    # Create table
    table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1]+1)
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Account / Period"
    hdr_cells[0].paragraphs[0].runs[0].bold = True
    
    # Header Shading
    tcPr = hdr_cells[0]._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E7E6E6')
    tcPr.append(shd)
    
    for i, col in enumerate(df.columns):
        date_str = str(col.date()) if hasattr(col, 'date') else str(col)
        hdr_cells[i+1].text = date_str
        hdr_cells[i+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hdr_cells[i+1].paragraphs[0].runs[0].bold = True
        
        # Header Shading
        tcPr = hdr_cells[i+1]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'E7E6E6')
        tcPr.append(shd)
        
    # Data Rows
    for idx, (index_name, row) in enumerate(df.iterrows()):
        row_cells = table.rows[idx+1].cells
        row_cells[0].text = str(index_name)
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        
        for i, val in enumerate(row):
            cell = row_cells[i+1]
            try:
                if isinstance(val, (int, float)):
                    cell.text = f"{val:,.0f}"
                else:
                    cell.text = str(val)
            except:
                cell.text = "-"
            
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            
    doc.add_paragraph("\n")

def save_to_word(ticker, markdown_content, filing_date, income, balance, cashflow, info=None):
    """
    Generate a professional Word report (Korean Brokerage Style).
    """
    doc = Document()
    
    # Set Font (Malgun Gothic for Korean support)
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    
    # --- Header Section ---
    # Company Name (Large)
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(f"{ticker} Analysis Report")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Sub-info line
    sub_p = doc.add_paragraph()
    sub_p.add_run(f"Filing Date: {filing_date} | ").bold = True
    sub_p.add_run(f"Generated by EDGARSTOCK-REPORT (Gemini)")
    sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph("-" * 50) # Divider
    
    # --- Key Metrics ---
    if info:
        create_key_metrics_table(doc, info)
    
    # --- 1. AI Analysis Section ---
    doc.add_heading('1. Comprehensive Analysis', level=1)
    
    # Simple Markdown to Word conversion
    clean_text = markdown_content.replace('**', '').replace('## ', '').replace('# ', '')
    
    for line in clean_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('Item') or line.startswith('Conclusion') or line.startswith('Risk') or line.startswith('Business'):
            doc.add_heading(line, level=2)
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
            
    doc.add_page_break()
    
    # --- 2. Financial Section ---
    doc.add_heading('2. Financial Statements (Key Metrics)', level=1)
    
    if income is not None:
        df_to_word_table(doc, income, "2-1. Income Statement")
        doc.add_page_break()
        
    if balance is not None:
        df_to_word_table(doc, balance, "2-2. Balance Sheet")
        doc.add_page_break()
        
    if cashflow is not None:
        df_to_word_table(doc, cashflow, "2-3. Cash Flow Statement")
    
    # --- Disclaimer ---
    doc.add_paragraph("-" * 50)
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run("Disclaimer: This report is generated by AI for informational purposes only. It is not financial advice. Please verify all information before making investment decisions.")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Save
    if not os.path.exists("reports"):
        os.makedirs("reports")
        
    filename = f"reports/{ticker}.docx"
    doc.save(filename)
    print(f"[{ticker}] Word report saved to {filename}")
    return filename

def save_to_file(ticker, content):
    """
    Save the report to a Markdown file.
    """
    if not os.path.exists("reports"):
        os.makedirs("reports")
        
    filename = f"reports/{ticker}.md"
    with open(filename, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"[{ticker}] Report saved to {filename}")

# --- Main Execution Block (Test) ---
if __name__ == "__main__":
    # Test with AAPL
    ticker = "AAPL"
    html, filing_date = get_sec_data(ticker)
    if html:
        print(f"[{ticker}] Successfully downloaded {len(html)} characters. (Date: {filing_date})")
        
        core_text = extract_sections(html)
        
        if core_text:
            report = analyze_with_gemini(core_text, ticker, filing_date)
            if report:
                save_to_file(ticker, report)
                print("Done!")
            else:
                print("Failed to generate report.")
        else:
            print("Failed to extract sections.")
        
    else:
        print(f"[{ticker}] Failed to download data.")
